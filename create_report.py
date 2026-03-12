# -*- coding: utf-8 -*-
"""
Create Word report: Program overview and exchange rules
for CHA Hospital Intern Turn Exchange System
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width = Cm(21.0)   # A4
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

# Heading styles
for level, (size, color_hex) in enumerate([
    (22, '1A5276'), (14, '1A5276'), (12, '2E86C1')
], start=1):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Malgun Gothic'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    pf = hs.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 14)
    pf.space_after = Pt(8)


def add_table(headers, rows, col_widths=None):
    """Add a styled table to the document"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A5276"/>')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, text in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(text)
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            # Alternate row shading
            if r % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF5FB"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    # Cell margins
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph()  # spacing after table
    return table


def add_bullet(text, bold_prefix=None, level=0):
    """Add a bullet point"""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        p.add_run(text)
    else:
        p.text = text
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    return p


def add_info_box(text, bg_color='D5E8F0'):
    """Add a highlighted info box using a single-cell table"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = text
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)
    # Set cell width to full page
    cell.width = Cm(16)
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after = Pt(4)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('차병원 인턴 턴표 교환 시스템')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
run.font.name = 'Malgun Gothic'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('프로그램 개요 및 교환 규칙 보고서')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
run.font.name = 'Malgun Gothic'

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('2026년 인턴 교육과정')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026.03')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════
doc.add_heading('목차', level=1)

toc_items = [
    '1. 시스템 개요',
    '2. 시스템 구성',
    '3. 교환 규칙',
    '   3.1 교환 불가 턴',
    '   3.2 필수과목 유지 규칙',
    '   3.3 분당 근무 최소 턴 규칙',
    '   3.4 휴가 교환 규칙',
    '   3.5 복합 교환 시 휴가 수지 균형',
    '4. 교환 유형',
    '   4.1 단일 교환',
    '   4.2 복합(체인) 교환',
    '5. 교환 장터',
    '6. 교환 시뮬레이션',
    '7. 관리자 기능',
    '8. 규칙 종합 요약표',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    if item.startswith('   '):
        p.paragraph_format.left_indent = Cm(1.5)
        p.runs[0].font.size = Pt(10)
    else:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(11)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ══════════════════════════════════════════════════════════════════
doc.add_heading('1. 시스템 개요', level=1)

doc.add_paragraph(
    '차병원 인턴 턴표 교환 시스템은 2026년 입사 인턴들이 '
    '자신의 근무 일정(턴)을 다른 인턴과 교환할 수 있도록 지원하는 '
    '웹 기반 애플리케이션입니다.'
)

doc.add_heading('시스템 기본 정보', level=2)

add_table(
    ['항목', '내용'],
    [
        ['시스템명', '차병원 인턴 턴표 교환소'],
        ['플랫폼', 'Streamlit (Python 웹 프레임워크)'],
        ['접속 URL', 'cha2026internchange.streamlit.app'],
        ['대상 사용자', '2026년 입사 인턴 전원 + 관리자(ADMIN)'],
        ['지원 기기', '모바일(기본) / PC 모두 지원, 전환 가능'],
        ['인증 방식', '이름 + 비밀번호 (초기 비밀번호: 1234)'],
        ['전체 턴 수', '13턴 (1턴~13턴)'],
        ['교환 가능 턴', '3턴~13턴 (11개 턴)'],
    ],
    col_widths=[4, 12]
)

doc.add_paragraph(
    '별도 앱 설치 없이 모바일 또는 PC 웹 브라우저에서 바로 접속하여 사용할 수 있으며, '
    '첫 로그인 시 비밀번호 변경을 안내합니다.'
)


# ══════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('2. 시스템 구성', level=1)

doc.add_heading('데이터 저장소', level=2)

doc.add_paragraph(
    '시스템은 Google Sheets를 주 데이터 저장소로 사용하며, '
    'gspread 라이브러리를 통해 실시간으로 연동됩니다. '
    '교환 시 구글 시트와 로컬 캐시(intern_data.json) 모두 동기화됩니다.'
)

add_table(
    ['시트 탭명', '용도'],
    [
        ['2026년 입사 인턴 스케쥴', '메인 스케줄 데이터 (인턴별 13턴 배치)'],
        ['2026년 입사 인턴 스케쥴_휴가', '휴가 배정 정보 (턴별 휴가 타입)'],
        ['비밀번호', '인턴별 비밀번호 관리'],
        ['교환이력', '교환 수락/거절 로그 기록'],
        ['장터', '교환 장터 게시물 관리'],
    ],
    col_widths=[6, 10]
)

doc.add_heading('과목 및 지역 체계', level=2)

doc.add_paragraph('인턴의 각 턴에는 과목(department)과 근무 지역이 배정됩니다.')

add_table(
    ['구분', '항목', '비고'],
    [
        ['필수과목', 'IM (내과)', '교환 후에도 1턴 이상 필수'],
        ['필수과목', 'GS (외과)', '교환 후에도 1턴 이상 필수'],
        ['필수과목', 'OB (산부인과)', '교환 후에도 1턴 이상 필수'],
        ['필수과목', 'PE (소아과)', '교환 후에도 1턴 이상 필수'],
        ['필수과목', '진로탐색', '교환 후에도 1턴 이상 필수'],
        ['근무 지역', '분당 (본원, 기본값)', '셀에 괄호 표시 없음 (예: IM)'],
        ['파견 지역', '일산', '셀에 괄호 표시 (예: IM(일산))'],
        ['파견 지역', '구미', '셀에 괄호 표시 (예: GS(구미))'],
        ['파견 지역', '강남', '셀에 괄호 표시 (예: IM(강남))'],
    ],
    col_widths=[3, 5, 8]
)

doc.add_heading('휴가 체계', level=2)

doc.add_paragraph('인턴에게는 1차, 2차 총 2회의 휴가가 배정됩니다.')

add_table(
    ['휴가', '대상 턴 범위', '비고'],
    [
        ['1차 휴가', '4턴, 5턴, 6턴, 7턴', '턴 번호 오름차순으로 첫 번째 휴가'],
        ['2차 휴가', '8턴, 9턴, 10턴, 11턴, 12턴, 13턴', '턴 번호 오름차순으로 두 번째 휴가'],
    ],
    col_widths=[3, 6, 7]
)

doc.add_paragraph(
    '휴가 타입은 알파벳-숫자 형식(예: A-4, B-2, C-3, D-1)으로 '
    '그룹과 순서를 나타내며, 교환 시 타입도 함께 교환됩니다.'
)


# ══════════════════════════════════════════════════════════════════
# 3. EXCHANGE RULES
# ══════════════════════════════════════════════════════════════════
doc.add_heading('3. 교환 규칙', level=1)

doc.add_paragraph(
    '모든 교환은 아래 규칙을 충족해야만 신청 및 실행이 가능합니다. '
    '규칙은 신청 시점과 수락 시점 모두에서 검증되며, '
    '수락 시점에는 최신 스케줄 데이터를 기반으로 재검증됩니다.'
)

# 3.1
doc.add_heading('3.1 교환 불가 턴', level=2)

add_info_box('1턴과 2턴은 교환 대상에서 제외됩니다. (고정 턴)', 'FADBD8')

doc.add_paragraph(
    '전체 13턴 중 1턴과 2턴은 교환이 불가능한 고정 턴으로 설정되어 있습니다. '
    '시뮬레이션, 장터, 교환 신청 등 모든 기능에서 1턴과 2턴은 후보에 포함되지 않습니다. '
    '따라서 실제 교환 가능한 턴은 3턴부터 13턴까지 총 11개 턴입니다.'
)

# 3.2
doc.add_heading('3.2 필수과목 유지 규칙', level=2)

add_info_box(
    '교환 후에도 IM, GS, OB, PE, 진로탐색 5개 필수과목이 '
    '각각 최소 1턴 이상 배정되어 있어야 합니다.',
    'D5F5E3'
)

doc.add_paragraph(
    '교환으로 인해 특정 필수과목이 스케줄에서 완전히 사라지는 경우, '
    '해당 교환은 차단됩니다. 검증은 교환에 관련된 모든 인턴(신청자, 상대방)에게 '
    '개별적으로 수행됩니다.'
)

p = doc.add_paragraph()
run = p.add_run('예시: ')
run.bold = True
p.add_run(
    '이규의 스케줄에 OB가 1턴만 있을 때, 해당 OB 턴을 다른 과목으로 교환하면 '
    'OB가 0턴이 되므로 교환이 차단됩니다.'
)

# 3.3
doc.add_heading('3.3 분당 근무 최소 턴 규칙', level=2)

add_info_box(
    '교환 후에도 분당(본원) 근무 턴 수가 최소 7개 이상이어야 합니다.',
    'D5F5E3'
)

doc.add_paragraph(
    '전체 13턴 중 분당(괄호 표시 없는 일반 과목)에 배치된 턴 수가 7개 미만이 되면 '
    '교환이 차단됩니다. 일산, 구미, 강남에 배치된 파견 턴은 분당 근무에 포함되지 않습니다.'
)

add_table(
    ['조건', '값'],
    [
        ['전체 턴 수', '13턴'],
        ['분당 최소 턴', '7턴'],
        ['최대 파견 턴', '6턴 (13 - 7)'],
    ],
    col_widths=[6, 10]
)

# 3.4
doc.add_heading('3.4 휴가 교환 규칙', level=2)

doc.add_paragraph('단일 턴 교환 시, 휴가 배정 여부에 따라 다음과 같이 적용됩니다:')

add_table(
    ['신청자 휴가 여부', '상대방 휴가 여부', '교환 가능 여부', '안내'],
    [
        ['O (휴가턴)', 'O (휴가턴)', '가능', '휴가 타입도 함께 교환됨'],
        ['O (휴가턴)', 'X (일반턴)', '불가', '복합 교환으로 상대에게서 다른 휴가턴을 받아오세요'],
        ['X (일반턴)', 'O (휴가턴)', '불가', '복합 교환으로 상대에게 다른 휴가턴을 주세요'],
        ['X (일반턴)', 'X (일반턴)', '가능', '휴가 규칙 적용 없음'],
    ],
    col_widths=[3.5, 3.5, 3, 6]
)

doc.add_paragraph(
    '한쪽만 휴가턴인 경우 단일 교환이 불가하며, '
    '복합 교환을 통해 휴가턴을 추가로 주고받아 균형을 맞춰야 합니다.'
)

# 3.5
doc.add_heading('3.5 복합 교환 시 휴가 수지 균형', level=2)

add_info_box(
    '복합 교환 시, 각 상대방별로 휴가턴을 준 횟수와 받은 횟수가 동일해야 합니다.',
    'D6EAF8'
)

doc.add_paragraph(
    '복합 교환(체인)에서 여러 턴을 동시에 교환할 때, '
    '상대방별로 휴가턴의 수지가 균형을 이루어야 합니다. '
    '전체 합산이 아닌 각 상대방과의 개별 균형이 요구됩니다.'
)

p = doc.add_paragraph()
run = p.add_run('예시: ')
run.bold = True
p.add_run(
    'A가 B에게 휴가턴 1개를 주었다면, B에게서도 휴가턴 1개를 받아와야 합니다. '
    'A가 C에게 휴가턴을 주고 B에게서 받는 것은 균형으로 인정되지 않습니다.'
)


# ══════════════════════════════════════════════════════════════════
# 4. EXCHANGE TYPES
# ══════════════════════════════════════════════════════════════════
doc.add_heading('4. 교환 유형', level=1)

# 4.1
doc.add_heading('4.1 단일 교환', level=2)

doc.add_paragraph(
    '신청자와 상대방이 1개 턴을 교환하는 기본 형태입니다.'
)

add_bullet('신청자가 상대방과 교환할 턴을 선택하여 요청을 보냅니다.')
add_bullet('신청 시점에 필수과목, 분당, 휴가 규칙이 모두 검증됩니다.')
add_bullet('상대방이 수락하면 최신 데이터로 재검증 후 즉시 교환이 실행됩니다.')
add_bullet('상대방이 거절하면 요청이 취소됩니다.')
add_bullet('재검증 실패 시(다른 교환으로 스케줄 변경됨) 자동 거절 처리됩니다.')

# 4.2
doc.add_heading('4.2 복합(체인) 교환', level=2)

doc.add_paragraph(
    '단일 교환으로 조건을 충족할 수 없는 경우, '
    '여러 상대방과 여러 턴을 동시에 교환하는 방식입니다. '
    '하나의 체인(chain_id)으로 묶여 관리됩니다.'
)

add_bullet('2~3개 턴을 동시에 교환 가능합니다.')
add_bullet('모든 상대방이 개별적으로 수락해야만 일괄 실행됩니다.', bold_prefix='전원 수락 필요: ')
add_bullet('한 명이라도 거절하면 체인 전체가 자동 취소됩니다.', bold_prefix='1인 거절 시 전체 취소: ')
add_bullet('모든 수락 완료 후에도 최신 데이터 기반 재검증이 실행됩니다.')

add_info_box(
    '복합 교환의 핵심 장점:\n'
    '단일 교환에서 한쪽만 휴가턴이어서 교환이 불가능한 경우,\n'
    '추가 턴을 묶어 휴가 수지 균형을 맞출 수 있습니다.',
    'D5F5E3'
)

p = doc.add_paragraph()
run = p.add_run('예시: ')
run.bold = True
p.add_run(
    'A의 4턴(휴가)을 B에게 주고 싶지만 단일 교환 불가 시, '
    '4턴 + 5턴(B의 휴가)을 함께 교환하면 '
    '양측 모두 휴가 1개씩 주고받아 균형이 맞으므로 교환이 가능해집니다.'
)


# ══════════════════════════════════════════════════════════════════
# 5. MARKETPLACE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('5. 교환 장터', level=1)

doc.add_paragraph(
    '교환 장터는 직접 상대방을 찾기 어려울 때 '
    '게시판 형태로 교환 희망 내용을 등록하는 기능입니다.'
)

doc.add_heading('등록 방식', level=2)

add_table(
    ['방식', '설명', '예시'],
    [
        ['줄 턴 지정', '특정 턴을 지정하여 교환 희망 등록', '5턴(GS)을 주고 IM을 받고 싶습니다'],
        ['아무 턴', '턴을 특정하지 않고 희망 과목만 등록', '아무 턴이나 주고 ANE를 받고 싶습니다'],
    ],
    col_widths=[3, 5, 8]
)

doc.add_heading('장터 기능', level=2)

add_bullet('장터 둘러보기: 다른 인턴의 게시물 확인 및 호환 가능 여부 자동 분석')
add_bullet('내 턴 올리기: 교환 희망 턴 등록 (받고 싶은 과목 지정 가능)')
add_bullet('내 게시물: 등록한 게시물 관리 (취소/마감 가능)')
add_bullet('교환 성사 시 해당 게시물이 자동으로 "완료" 처리됩니다.')
add_bullet('같은 사람이 같은 턴으로 중복 등록은 차단됩니다.')


# ══════════════════════════════════════════════════════════════════
# 6. SIMULATION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('6. 교환 시뮬레이션', level=1)

doc.add_paragraph(
    '교환 가능 여부를 사전에 탐색할 수 있는 시뮬레이션 기능을 제공합니다. '
    '총 3가지 모드가 있습니다.'
)

add_table(
    ['모드', '설명', '활용 예시'],
    [
        ['특정 턴 교환',
         '선택한 턴을 교환할 수 있는 모든 파트너 목록 표시',
         '3턴을 교환하고 싶은데 누구와 가능한지 확인'],
        ['특정 과목 받기',
         '원하는 과목을 받을 수 있는 모든 (파트너, 턴) 조합 탐색',
         'ANE를 받고 싶은데 어떤 교환이 가능한지 확인'],
        ['복합 교환 탐색',
         '2~3개 턴 동시 교환의 유효한 조합 자동 탐색\n(단독 불가 조합만 필터 가능)',
         '단일 교환 불가 시 복합으로 가능한 조합 확인'],
    ],
    col_widths=[3, 6, 7]
)

doc.add_paragraph(
    '시뮬레이션 결과에서 가능한 파트너에게 바로 교환 요청을 보낼 수 있습니다. '
    '복합 교환 탐색 시 결과 상한은 200개이며, '
    '3개 조합은 후보 수 200개 이하일 때만 탐색됩니다.'
)


# ══════════════════════════════════════════════════════════════════
# 7. ADMIN
# ══════════════════════════════════════════════════════════════════
doc.add_heading('7. 관리자 기능', level=1)

doc.add_paragraph(
    'ADMIN 계정으로 로그인 시 관리자 전용 대시보드에 접근할 수 있습니다. '
    '총 6개 탭으로 구성됩니다.'
)

add_table(
    ['탭', '기능', '주요 내용'],
    [
        ['스케줄 통계', '턴별/과목별 배치 현황 분석',
         '과목 분포, 지역별 인원, 필수과목 충족 현황'],
        ['전체 스케줄', '전체 인턴 스케줄 조회',
         '인턴별 상세 조회, 휴가 음영 표시'],
        ['휴가 현황', '휴가 배정 상태 관리',
         '턴별 휴가 인원, 타입별 분포, 배정 현황'],
        ['교환 이력', '교환 수락/거절 로그',
         '전체 이력 최신순 조회, 대기 중 요청'],
        ['장터 현황', '장터 게시물 관리',
         '상태별 필터링, 활성/마감/완료 현황'],
        ['비밀번호 관리', '인턴 비밀번호 관리',
         '비밀번호 초기화, 변경 상태 확인'],
    ],
    col_widths=[3, 5, 8]
)


# ══════════════════════════════════════════════════════════════════
# 8. SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('8. 규칙 종합 요약표', level=1)

doc.add_paragraph('아래 표는 시스템의 모든 교환 규칙을 종합 정리한 것입니다.')

summary_table = add_table(
    ['규칙', '조건', '적용 시점'],
    [
        ['교환 불가 턴', '1턴, 2턴 교환 절대 불가', '모든 교환'],
        ['필수과목 유지',
         'IM, GS, OB, PE, 진로탐색 각 1턴 이상',
         '신청 시 + 수락 시 재검증'],
        ['분당 최소 7턴',
         '전체 13턴 중 분당 배치 7턴 이상',
         '신청 시 + 수락 시 재검증'],
        ['단일 휴가 교환',
         '한쪽만 휴가턴이면 차단 (둘 다 휴가면 OK)',
         '단일 교환'],
        ['복합 휴가 수지',
         '상대방별 휴가턴 주고받기 균형 (given == received)',
         '복합/체인 교환'],
        ['중복 요청 방지',
         '동일 sender-receiver-turn의 pending 요청 존재 시 차단',
         '교환 신청'],
        ['수락 시 재검증',
         '수락 전 시트에서 최신 데이터 갱신 후 전체 규칙 재검증',
         '수락 처리'],
        ['체인 전체 거절',
         '복합 교환에서 1명이라도 거절 시 전체 자동 취소',
         '체인 거절 시'],
        ['장터 중복 방지',
         '같은 인턴이 같은 턴으로 활성 게시물 중복 등록 차단',
         '장터 등록'],
    ],
    col_widths=[3.5, 7, 5.5]
)

doc.add_paragraph()

# Final note
add_info_box(
    '본 보고서는 차병원 인턴 턴표 교환 시스템(app.py)의 소스 코드를 기반으로 작성되었습니다.\n'
    '시스템 접속: cha2026internchange.streamlit.app',
    'EBF5FB'
)


# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'report_exchange_rules.docx')
doc.save(output_path)
print(f'Report saved: {output_path}')
